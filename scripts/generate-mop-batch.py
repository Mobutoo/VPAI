#!/usr/bin/env python3
"""
generate-mop-batch.py — Génère les 5 MOP procédures OTN (M0.20–M0.24) au format DOCX.
Source : mop-incidents-alignment-review.xlsx (onglet "Révision enrichie")
Output : docs/mop/OPSSD-03.03.M0.XX-P-V1.docx (5 fichiers)
"""

import os
import sys
import re
from datetime import datetime

import openpyxl
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ─── Chemins ─────────────────────────────────────────────────────────────────
BASE_DIR = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
XLSX_PATH = os.path.join(BASE_DIR, "docs", "mop-incidents-alignment-review.xlsx")
FALLBACK_XLSX = "/tmp/mop-align-review.xlsx"
OUT_DIR = os.path.join(BASE_DIR, "docs", "mop")

# ─── Couleurs VPAI ───────────────────────────────────────────────────────────
BLEU_MARINE = RGBColor(0x1F, 0x38, 0x64)   # #1F3864
BLEU_CLAIR  = RGBColor(0xBD, 0xD7, 0xEE)   # #BDD7EE
GRIS_CLAIR  = RGBColor(0xF2, 0xF2, 0xF2)   # #F2F2F2
ORANGE      = RGBColor(0xF4, 0xB9, 0x42)   # #F4B942
VERT        = RGBColor(0x70, 0xAD, 0x47)   # #70AD47 (MINEUR)
ROUGE       = RGBColor(0xC0, 0x00, 0x00)   # #C00000 (CRITIQUE)


def fmt_date(val):
    """datetime → DD/MM/YYYY, string passthrough."""
    if isinstance(val, datetime):
        return val.strftime("%d/%m/%Y")
    return str(val) if val else ""


def add_colored_heading(doc, text, level, bg_color=None, text_color=None):
    """Ajoute un heading stylisé."""
    p = doc.add_paragraph()
    p.style = f"Heading {level}"
    run = p.add_run(text)
    if text_color:
        run.font.color.rgb = text_color
    if bg_color:
        pPr = p._p.get_or_add_pPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), rgb_hex(bg_color))
        pPr.append(shd)
    return p


def add_shaded_para(doc, text, bg_color, bold=False, font_size=10):
    """Paragraphe avec fond coloré."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(font_size)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), rgb_hex(bg_color))
    pPr.append(shd)
    return p


def rgb_hex(color: RGBColor) -> str:
    return f"{color[0]:02X}{color[1]:02X}{color[2]:02X}"


def set_cell_bg(cell, color: RGBColor):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), rgb_hex(color))
    tcPr.append(shd)


def set_cell_text(cell, text, bold=False, font_size=10, color=None):
    cell.text = ""
    para = cell.paragraphs[0]
    run = para.add_run(str(text) if text else "")
    run.bold = bold
    run.font.size = Pt(font_size)
    if color:
        run.font.color.rgb = color
    return run


def add_multiline_content(doc, content, font_size=10, bullet=False):
    """Ajoute un bloc texte multi-lignes. Lignes vides = séparateur."""
    if not content:
        doc.add_paragraph()
        return
    lines = str(content).split("\n")
    for line in lines:
        line = line.strip()
        if not line:
            continue
        if bullet or line.startswith(("- ", "• ")):
            p = doc.add_paragraph(style="List Bullet")
            run = p.add_run(line.lstrip("- •").strip())
        elif re.match(r"^\d+\.", line):
            p = doc.add_paragraph(style="List Number")
            run = p.add_run(line[line.index(".")+1:].strip())
        else:
            p = doc.add_paragraph()
            run = p.add_run(line)
        run.font.size = Pt(font_size)


def build_cover_table(doc, mop):
    """Cartouche document (tableau 2 colonnes)."""
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Ligne titre
    row = table.rows[0]
    c0, c1 = row.cells[0], row.cells[1]
    c0.merge(c1)
    set_cell_bg(c0, BLEU_MARINE)
    p = c0.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("MÉTHODE D'OPÉRATION ET DE PROCÉDURE")
    run.bold = True
    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    run.font.size = Pt(12)

    def add_row(label, value, label_bg=BLEU_CLAIR):
        row = table.add_row()
        c0, c1 = row.cells[0], row.cells[1]
        set_cell_bg(c0, label_bg)
        set_cell_text(c0, label, bold=True, font_size=9)
        set_cell_text(c1, value, font_size=9)

    add_row("Référence", mop["doc_reference"])
    add_row("Titre", mop["doc_title"])
    add_row("Type", "MOP")
    add_row("Confidentialité", "Interne")
    add_row("Propriétaire", "Service Technique")
    add_row("Vérificateur", mop["doc_verificateur"])
    add_row("Approbateur", "Service ISO")
    add_row("Date d'entrée en vigueur", fmt_date(mop["doc_date_vigueur"]))
    add_row("Date de révision planifiée", fmt_date(mop["doc_date_revision"]))


def build_revision_table(doc, mop):
    """Tableau des révisions."""
    p = doc.add_paragraph()
    run = p.add_run("HISTORIQUE DES RÉVISIONS")
    run.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = BLEU_MARINE

    table = doc.add_table(rows=2, cols=5)
    table.style = "Table Grid"

    headers = ["Rev.", "Date", "Préparé par", "Vérifié par", "Approuvé par"]
    hrow = table.rows[0]
    for i, h in enumerate(headers):
        set_cell_bg(hrow.cells[i], BLEU_MARINE)
        set_cell_text(hrow.cells[i], h, bold=True, font_size=9, color=RGBColor(0xFF, 0xFF, 0xFF))

    drow = table.rows[1]
    vals = ["1.0", fmt_date(mop["rev1_date"]), mop["rev1_prepared_by"],
            mop["rev1_verified_by"], "Service ISO"]
    for i, v in enumerate(vals):
        set_cell_text(drow.cells[i], v, font_size=9)


def build_equipements(doc, mop):
    """Bloc équipements couverts."""
    add_shaded_para(doc, "ÉQUIPEMENTS COUVERTS", BLEU_CLAIR, bold=True, font_size=9)
    add_multiline_content(doc, mop["equipements_couverts"], font_size=9)


def build_contacts(doc, mop):
    """Bloc contacts escalade."""
    add_shaded_para(doc, "CONTACTS D'ESCALADE", BLEU_CLAIR, bold=True, font_size=9)
    add_multiline_content(doc, mop["contacts_escalade"], font_size=9)


def build_section(doc, num, title, content):
    """Section standard avec numéro."""
    p = doc.add_paragraph()
    p.style = "Heading 1"
    run = p.add_run(f"Section {num} — {title}")
    run.font.color.rgb = BLEU_MARINE

    if content:
        add_multiline_content(doc, content)


def build_section3(doc, mop):
    """Section 3 — Procédure d'intervention (phases 3.1.1–3.1.4)."""
    p = doc.add_paragraph()
    p.style = "Heading 1"
    run = p.add_run("Section 3 — Description du processus (comment)")
    run.font.color.rgb = BLEU_MARINE

    # REX contexte
    if mop.get("s3_content"):
        p_rex = doc.add_paragraph()
        p_rex.add_run("Contexte terrain : ").bold = True
        p_rex.add_run(mop["s3_content"]).font.size = Pt(10)

    # 3.1 — Phases
    p31 = doc.add_paragraph()
    p31.style = "Heading 2"
    p31.add_run("3.1 — Phases d'intervention")

    phases = [
        ("3.1.1", "Phase 1 — Qualification initiale", "phase1_content"),
        ("3.1.2", "Phase 2 — Diagnostic et actions", "phase2_content"),
        ("3.1.3", "Phase 3 — Résolution et restoration", "phase3_content"),
        ("3.1.4", "Phase 4 — Clôture et retour à la normale", "phase4_content"),
    ]
    for num, title, key in phases:
        p_ph = doc.add_paragraph()
        p_ph.style = "Heading 3"
        run = p_ph.add_run(f"{num} {title}")

        content = mop.get(key, "")
        if content:
            add_multiline_content(doc, content)


def generate_mop_docx(mop: dict, out_path: str):
    """Génère un fichier DOCX pour un MOP famille."""
    doc = Document()

    # ── Styles de base ──────────────────────────────────────────────────────
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10)

    for h_level in (1, 2, 3):
        try:
            hs = doc.styles[f"Heading {h_level}"]
            hs.font.name = "Calibri"
            hs.font.bold = True
            if h_level == 1:
                hs.font.size = Pt(12)
            elif h_level == 2:
                hs.font.size = Pt(11)
            else:
                hs.font.size = Pt(10)
        except KeyError:
            pass

    # ── Sections ─────────────────────────────────────────────────────────────
    # Page de garde
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_p.add_run(mop["doc_title"])
    title_run.bold = True
    title_run.font.size = Pt(16)
    title_run.font.color.rgb = BLEU_MARINE

    ref_p = doc.add_paragraph()
    ref_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    ref_p.add_run(mop["doc_reference"]).font.size = Pt(11)

    doc.add_paragraph()

    # Cartouche
    build_cover_table(doc, mop)
    doc.add_paragraph()

    # Révisions
    build_revision_table(doc, mop)
    doc.add_paragraph()

    # Résumé
    add_shaded_para(doc, "RÉSUMÉ", BLEU_CLAIR, bold=True, font_size=9)
    rsum_p = doc.add_paragraph()
    rsum_p.add_run(mop.get("doc_resume", "")).font.size = Pt(10)
    doc.add_paragraph()

    # Équipements + contacts
    build_equipements(doc, mop)
    doc.add_paragraph()
    build_contacts(doc, mop)

    # Page break avant sections procédure
    doc.add_page_break()

    # Section 1 — Objet
    build_section(doc, 1, "Objet (quoi / pourquoi)", mop.get("s1_content", ""))
    doc.add_paragraph()

    # Section 2 — Rôles
    build_section(doc, 2, "Rôles et responsabilités (qui)", mop.get("s2_content", ""))
    doc.add_paragraph()

    # Section 3 — Processus + phases
    build_section3(doc, mop)
    doc.add_paragraph()

    # Section 4 — Glossaire
    build_section(doc, 4, "Glossaire", mop.get("s4_content", ""))

    doc.save(out_path)
    print(f"  → {out_path} ({os.path.getsize(out_path):,} B)")


def extract_mop_data(xlsx_path: str) -> list:
    """Extrait les 5 MOP depuis l'onglet 'Révision enrichie'."""
    wb = openpyxl.load_workbook(xlsx_path)
    ws = wb["Révision enrichie"]

    # Colonnes : C=SUPERVISION, D=PHOTONIQUE, E=CARTE, F=CHASSIS, G=PORT NNI/UNI
    col_idx = [3, 4, 5, 6, 7]  # 1-based: C=3, D=4, E=5, F=6, G=7

    field_map = {
        "doc_reference":       5,
        "doc_title":           6,
        "doc_resume":          7,
        "doc_verificateur":    8,
        "doc_date_vigueur":    9,
        "doc_date_revision":   10,
        "rev1_date":           12,
        "rev1_prepared_by":    13,
        "rev1_verified_by":    14,
        "equipements_couverts": 16,
        "contacts_escalade":   18,
        "s1_content":          20,
        "s2_content":          22,
        "s3_content":          24,
        "phase1_content":      25,
        "phase2_content":      26,
        "phase3_content":      27,
        "phase4_content":      28,
        "s4_content":          30,
        "sources_kb":          32,
    }

    mops = []
    for ci in col_idx:
        mop = {}
        for field, row in field_map.items():
            val = ws.cell(row=row, column=ci).value
            mop[field] = val if val is not None else ""
        # Normalise la référence : ajoute -V1 si absent
        ref = str(mop["doc_reference"])
        if ref and not ref.endswith("-V1"):
            ref = ref + "-V1"
        mop["doc_reference"] = ref
        mops.append(mop)

    return mops


def main():
    # Trouver le fichier source
    xlsx_path = XLSX_PATH if os.path.exists(XLSX_PATH) else FALLBACK_XLSX
    if not os.path.exists(xlsx_path):
        print(f"ERREUR : fichier source introuvable : {xlsx_path}", file=sys.stderr)
        sys.exit(1)

    print(f"Source : {xlsx_path}")
    os.makedirs(OUT_DIR, exist_ok=True)

    mops = extract_mop_data(xlsx_path)
    print(f"  {len(mops)} MOP extraits de 'Révision enrichie'")

    for mop in mops:
        ref = mop["doc_reference"]
        out_file = os.path.join(OUT_DIR, f"{ref}.docx")
        print(f"\n[{ref}] {mop['doc_title']}")
        generate_mop_docx(mop, out_file)

    print(f"\n✓ {len(mops)} fichiers générés dans {OUT_DIR}/")


if __name__ == "__main__":
    main()
